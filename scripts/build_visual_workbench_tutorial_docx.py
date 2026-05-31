from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "英国驿站仓储物流工作台图文操作教程-最终版.docx"
SHOT_DIR = ROOT / "deliverables" / "tutorial-screenshots"


BLUE = "0E7490"
DARK = "0F172A"
LIGHT_BLUE = "E8F7FB"
LIGHT_GRAY = "F8FAFC"
BORDER = "CBD5E1"
AMBER = "FEF3C7"


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=100, start=120, bottom=100, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    values = {"top": top, "start": start, "bottom": bottom, "end": end}
    for key, value in values.items():
        node = tc_mar.find(qn(f"w:{key}"))
        if node is None:
            node = OxmlElement(f"w:{key}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    grid = table._tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        table._tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for index, cell in enumerate(row.cells):
            cell.width = width_to_inches(widths[index])
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[index]))
            tc_w.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def width_to_inches(dxa: int):
    return Inches(dxa / 1440)


def set_run(run, size=10.5, bold=False, color=DARK):
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_para(doc, text="", size=10.5, bold=False, color=DARK, align=None, space_after=6):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(space_after)
    paragraph.paragraph_format.line_spacing = 1.25
    if align is not None:
        paragraph.alignment = align
    run = paragraph.add_run(text)
    set_run(run, size=size, bold=bold, color=color)
    return paragraph


def add_heading(doc, text, level=1):
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(12 if level == 1 else 8)
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run(text)
    set_run(run, size=17 if level == 1 else 13, bold=True, color=BLUE if level == 1 else DARK)
    return paragraph


def add_caption(doc, text):
    return add_para(doc, text, size=9.2, color="475569", align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)


def add_note(doc, title, body, fill=LIGHT_BLUE):
    table = doc.add_table(rows=1, cols=1)
    set_table_width(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "BAE6FD")
    paragraph = cell.paragraphs[0]
    paragraph.paragraph_format.space_after = Pt(3)
    run = paragraph.add_run(title)
    set_run(run, size=10.5, bold=True, color=BLUE)
    paragraph = cell.add_paragraph()
    paragraph.paragraph_format.space_after = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.25
    run = paragraph.add_run(body)
    set_run(run, size=10, color=DARK)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def add_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    widths = widths or [int(9360 / len(headers))] * len(headers)
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.cell(0, i)
        set_cell_shading(cell, "E2E8F0")
        paragraph = cell.paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = paragraph.add_run(header)
        set_run(run, size=9.5, bold=True, color="1E3A5F")
    for row_data in rows:
        row = table.add_row()
        for i, value in enumerate(row_data):
            cell = row.cells[i]
            set_cell_shading(cell, "FFFFFF")
            paragraph = cell.paragraphs[0]
            paragraph.paragraph_format.space_after = Pt(0)
            paragraph.paragraph_format.line_spacing = 1.2
            run = paragraph.add_run(str(value))
            set_run(run, size=9.2, bold=(i == 0), color=DARK if i != 0 else "0F172A")
    set_table_width(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)
    return table


def add_image(doc, filename, title, note=None):
    path = SHOT_DIR / filename
    if not path.exists():
        add_note(doc, "截图缺失", f"未找到截图文件：{path}", fill=AMBER)
        return
    title_paragraph = add_para(doc, title, size=10.5, bold=True, color=DARK, space_after=4)
    title_paragraph.paragraph_format.keep_with_next = True
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.keep_with_next = True
    run = paragraph.add_run()
    run.add_picture(str(path), width=Inches(7.0))
    paragraph.paragraph_format.space_after = Pt(2)
    if note:
        add_caption(doc, note)


def setup_document():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(0.58)
    section.bottom_margin = Inches(0.55)
    section.left_margin = Inches(0.62)
    section.right_margin = Inches(0.62)

    styles = doc.styles
    styles["Normal"].font.name = "Microsoft YaHei"
    styles["Normal"]._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    styles["Normal"].font.size = Pt(10.5)

    for section in doc.sections:
        header = section.header.paragraphs[0]
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = header.add_run("英国驿站仓储/物流工作台图文操作教程")
        set_run(run, size=8.5, color="64748B")
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer.add_run("内部培训资料 - 以系统最新页面为准")
        set_run(run, size=8.5, color="64748B")
    return doc


def cover(doc):
    for _ in range(3):
        doc.add_paragraph()
    add_para(doc, "操作手册", size=13, bold=True, color=BLUE)
    add_para(doc, "英国驿站仓储/物流工作台", size=26, bold=True, color=DARK, space_after=2)
    add_para(doc, "图文操作教程 - 最终版", size=20, bold=True, color=DARK, space_after=14)
    add_para(doc, "适用于客户、运营、仓库员工、财务和管理者。目标是让新人看完后知道每个入口点哪里、每一步做什么、出现异常时找谁处理。", size=11, color="334155")

    add_table(
        doc,
        ["项目", "说明"],
        [
            ["适用系统", "客户工作台、运营后台、仓库作业台"],
            ["教程范围", "登录注册、客户提需求/入库/出库/账号资料、运营审核、物流/账单、仓库扫码收货/拣货/库位"],
            ["阅读方式", "先看流程总览，再按自己的角色阅读对应章节；看到截图编号后，对照下方说明操作。"],
            ["版本日期", "2026-05-27"],
        ],
        widths=[1900, 7460],
    )
    add_note(doc, "一句话理解", "客户在工作台提交需求、入库、SKU、出库、退货和账单动作；运营在后台审核和推进；仓库在作业台扫码、收货、上架、拣货、打包和交运；老板看待办、异常、账单和整体履约状态。")
    doc.add_page_break()


def toc_and_map(doc):
    add_heading(doc, "1. 先看整体流程", 1)
    add_table(
        doc,
        ["角色", "每天主要看哪里", "最常做什么", "结果会体现在哪里"],
        [
            ["客户", "客户工作台 /portal", "提交需求、入库预报、维护 SKU、创建出库、查看账单和异常", "自己的待办、状态卡片、查进度、费用账单"],
            ["运营", "运营后台 /ops", "审核资料、报价、推进入库/出库/物流异常、生成账单", "运营待办中心、各业务队列、客户侧状态同步"],
            ["仓库", "仓库作业台 /warehouse", "扫码收货、上架、拣货、打包复核、交运、登记异常", "库存流水、作业状态、客户出库/入库进度"],
            ["老板/管理者", "运营后台总览 + 客户/仓库状态", "看异常、看效率、看账单、看客户状态和权限", "总览指标、异常队列、账单与权限审计"],
        ],
        widths=[1300, 2300, 3400, 2360],
    )
    add_note(doc, "操作顺序建议", "新业务先从客户提交需求开始；已确定发货时先做 SKU 和入库预报；货到仓后由仓库扫码收货并上架；客户创建出库后，仓库拣货打包，运营处理面单、追踪号和账单。")


def customer_sections(doc):
    add_heading(doc, "2. 客户工作台怎么用", 1)
    add_image(doc, "01-login.png", "图 1：客户登录和注册入口", "新客户先注册，老客户直接登录。登录后进入客户工作台。")
    add_table(
        doc,
        ["编号", "点哪里", "做什么", "注意事项"],
        [
            ["1", "手机 / 邮箱输入框", "输入注册时使用的手机号或邮箱。", "客户账号不用客服邀请码，注册后自动生成客户编号。"],
            ["2", "密码输入框", "输入客户自己设置的密码。", "忘记密码时走重置密码流程。"],
            ["3", "登录客户工作台", "提交登录，成功后进入客户工作台首页。", "如果已登录，再点用户工作台会直接进入首页。"],
            ["4", "免费注册", "新客户创建自己的工作台账号。", "注册后继续完善公司、VAT、EORI、平台店铺信息。"],
        ],
        widths=[800, 2100, 3300, 3160],
    )

    add_image(doc, "02-customer-portal.png", "图 2：客户工作台首页", "客户每天先看账号状态、下一步建议、功能入口和待办。")
    add_table(
        doc,
        ["编号", "入口", "用途"],
        [
            ["1", "我的待办", "显示需要客户处理的事项，例如确认报价、补资料、确认账单、处理物流异常。"],
            ["2", "提交需求", "还没有明确入库/出库前，先提交平台、货量、SKU 和服务需求，让运营评估报价。"],
            ["3", "入库预报", "货准备发到英国仓前创建 ASN，填写箱数、SKU、追踪号和资料。"],
            ["4", "SKU 档案", "维护商品编码、品名、条码、分类、库存预警，后续入库/出库都按 SKU 识别。"],
            ["5", "出库申请", "创建一件代发、FBA 补货或其他出库需求。"],
            ["6", "费用账单", "查看仓租、操作费、物流费、退货质检费等费用，并确认或提出争议。"],
            ["7", "账号资料", "维护公司资料、VAT、EORI、平台店铺和密码；未认证时优先处理这里。"],
        ],
        widths=[700, 1900, 6760],
    )

    add_image(doc, "03-customer-inbound.png", "图 3：创建 ASN 入库预报", "客户在货到仓前填写这页，方便仓库提前匹配货物。")
    add_table(
        doc,
        ["编号", "字段/按钮", "怎么填"],
        [
            ["1", "返回客户工作台", "不提交时返回首页，检查其他待办。"],
            ["2", "客户 / 店铺名称", "填写公司名或平台店铺名，方便运营和仓库识别。"],
            ["3", "箱数 / 托数", "填写预计到仓数量，仓库会按箱数验货。"],
            ["4", "SKU 明细", "按 SKU、品名、数量、箱数整理，建议一行一个 SKU。"],
            ["5", "提交入库预报", "提交后运营和仓库会看到任务；资料缺失时客户工作台会出现补资料待办。"],
        ],
        widths=[700, 2200, 6460],
    )

    add_image(doc, "04-customer-outbound.png", "图 4：创建出库申请", "客户创建发货需求，系统会预占库存，运营和仓库再接着处理。")
    add_table(
        doc,
        ["编号", "字段/按钮", "怎么填"],
        [
            ["1", "返回客户工作台", "返回首页，不影响已保存/已提交的记录。"],
            ["2", "尾程渠道", "选择 Royal Mail、DPD、Parcelforce 等渠道；后续可由运营按规则匹配。"],
            ["3", "订单数/收件人", "填写订单数量、收件人和期望发货日。"],
            ["4", "出库 SKU 明细", "一行一个 SKU 和数量，系统用于库存预占和拣货。"],
            ["5", "提交出库申请", "提交后进入运营/仓库队列，仓库按出库单拣货、复核、打包和交运。"],
        ],
        widths=[700, 2200, 6460],
    )

    add_image(doc, "05-customer-account.png", "图 5：账号与公司资料", "客户资料越完整，运营审核和正式入库越顺。")
    add_table(
        doc,
        ["编号", "模块", "说明"],
        [
            ["1", "账号状态", "显示未认证、已认证或暂停；未认证时仍可提交需求和维护资料，但正式账期/合同价会受限制。"],
            ["2", "公司资料", "维护公司名称、联系人、手机号、邮箱、VAT、EORI、公司地址。"],
            ["3", "平台 / 店铺", "填写 Amazon、TikTok Shop、Shopify、eBay 等平台和店铺链接。"],
            ["4", "修改密码", "客户自主修改密码；员工不能代客户保存密码。"],
        ],
        widths=[700, 2200, 6460],
    )


def ops_sections(doc):
    add_heading(doc, "3. 运营后台怎么用", 1)
    add_image(doc, "06-ops-overview.png", "图 6：运营后台总览", "运营每天从左侧模块和中间待办进入，先处理异常，再推进正常队列。")
    add_table(
        doc,
        ["编号", "入口", "运营要做什么"],
        [
            ["1", "总览", "看今天的客户待办、异常优先队列、状态汇总和业务提醒。"],
            ["2", "询盘", "处理客户需求，补充报价草案、报价有效期、客服备注，并推进客户确认。"],
            ["3", "入库", "审核 ASN、装箱单、SKU 清单、追踪号和到仓信息，缺资料时标记给客户。"],
            ["4", "库存", "看库存、库位、批次、冻结、残次、预警和调整审批。"],
            ["5", "出库", "处理客户出库申请、批量波次、拣货单、面单和发货状态。"],
            ["6", "物流", "处理承运商匹配、运费规则、追踪号回传、派送异常和赔付资料。"],
            ["7", "账单", "生成费用、锁定月结、处理客户确认/争议/付款核销。"],
        ],
        widths=[700, 1700, 6960],
    )

    add_image(doc, "07-ops-logistics.png", "图 7：运营扩展区、权限和自助配置", "这一屏主要给运营负责人、管理员和老板查看上线体检、权限、报表与客户自助项。")
    add_table(
        doc,
        ["编号", "区域", "用途"],
        [
            ["1", "上线体检", "上线前检查客户、账单、资料、仓储物流闭环是否完整。"],
            ["2", "客户自助工单处理队列", "处理客户提交的物流异常、库存调整、账单争议、退货售后和资料补充。"],
            ["3", "高级筛选/保存视图", "按客户、时间、仓库、平台、状态筛选，并保存常用运营看板。"],
            ["4", "权限矩阵和敏感操作", "配置角色可访问模块，账单锁定、库存调整等敏感动作需要二次确认。"],
        ],
        widths=[700, 2500, 6160],
    )
    add_note(doc, "老板重点看什么", "老板不需要逐条操作每个订单，重点看：待处理数量是否堆积、异常是否有人跟进、客户账单是否锁定、库存调整是否有审批、仓库出库/入库是否按 SLA 推进。")


def warehouse_sections(doc):
    add_heading(doc, "4. 仓库作业台怎么用", 1)
    add_image(doc, "08-warehouse-workbench.png", "图 8：仓库任务工作台", "仓库员工每天从这里扫码、找任务、推进收货、上架、拣货、打包和交运。")
    add_table(
        doc,
        ["编号", "入口/区域", "用途"],
        [
            ["1", "作业台", "选择当前作业模式，例如拣货、收货、上架、复核。"],
            ["2", "入库收货", "按 ASN 验货，核对箱数、SKU、包装单和外箱标签。"],
            ["3", "出库拣货", "按出库单或波次拣货，逐个扫描 SKU 复核。"],
            ["4", "库位管理", "维护库区、货架、库位、容量和是否允许混放。"],
            ["5", "打印", "打印拣货单、临时面单、补打面单或批量导出作业资料。"],
        ],
        widths=[700, 2200, 6460],
    )

    add_image(doc, "09-warehouse-scanning.png", "图 9：扫码后的系统体现", "扫码枪相当于键盘输入，扫完通常自动回车；系统根据码内容定位任务并更新进度。")
    add_table(
        doc,
        ["编号", "扫什么", "系统会怎么显示"],
        [
            ["1", "扫码输入框", "光标放在这里后，用扫码枪扫 ASN、出库单、SKU、追踪号或库位码。"],
            ["2", "ASN / 入库单", "定位入库任务，显示客户、箱数、SKU、资料状态和当前收货步骤。"],
            ["3", "SKU 条码", "确认当前商品是否属于任务；匹配则增加复核进度，不匹配则提示异常。"],
            ["4", "库位码", "确认商品上架或拣货的位置，系统写入库存流水和库位余额。"],
            ["5", "异常登记", "发现缺货、错货、破损、面单错误、截单等问题时登记，运营和老板能在异常队列看到。"],
        ],
        widths=[700, 2100, 6560],
    )
    add_note(doc, "扫码硬件怎么接", "常见 USB 或蓝牙扫码枪接入电脑后会被识别成键盘。员工只要点进扫码输入框，扫码枪扫出的条码会自动填入；如果扫码枪设置了自动回车，系统会立即执行确认扫描。")


def flow_sections(doc):
    add_heading(doc, "5. 常用业务流程", 1)
    add_table(
        doc,
        ["流程", "客户做什么", "运营做什么", "仓库做什么", "老板看什么"],
        [
            ["新客户开通", "免费注册，完善公司/VAT/EORI/平台店铺。", "审核客户资料，更新认证状态。", "暂不操作。", "看客户是否已认证、资料是否完整。"],
            ["报价需求", "点提交需求，写平台、货量、SKU 和服务要求。", "评估报价，写报价草案和有效期。", "必要时评估库位/作业难度。", "看高价值线索和报价待确认数量。"],
            ["入库", "点入库预报，提交箱数、SKU、追踪号和资料。", "审核资料，缺资料时标记待客户补充。", "扫码收货、登记差异、上架到库位。", "看到入库 SLA、异常和库存是否生成。"],
            ["出库", "点出库申请，提交 SKU、数量、渠道和地址。", "匹配渠道、生成面单、处理费用。", "扫码拣货、复核、打包、交运。", "看出库是否积压、物流是否回传。"],
            ["账单", "查看费用，确认账单或发起争议。", "生成费用、锁定月结、核销付款。", "提供作业依据。", "看应收、争议、已付和未结算。"],
        ],
        widths=[1200, 2200, 2200, 1900, 1860],
    )

    add_heading(doc, "6. 常见问题", 1)
    add_table(
        doc,
        ["问题", "怎么处理"],
        [
            ["客户说看不到正式库存", "先确认入库是否已由仓库推进到已上架；未上架前只会显示预报或待处理状态。"],
            ["客户提交出库后没进度", "运营先看出库队列是否待审核或缺库存；仓库看是否已生成拣货任务。"],
            ["扫码没反应", "确认光标在扫码输入框；用记事本测试扫码枪是否能输出条码；再确认扫码枪是否自动回车。"],
            ["扫码提示不匹配", "先核对当前任务、SKU、库位是否正确；如果实物和系统不一致，登记异常，不要强行推进。"],
            ["账单金额有争议", "客户在账单里发起争议；运营查看费用来源单据，必要时财务复核后更新状态。"],
            ["员工无法登录后台", "确认员工是否在白名单内；运营/仓库/管理员账号权限不同，不能共用客户账号。"],
        ],
        widths=[2600, 6760],
    )

    add_heading(doc, "7. 最终检查清单", 1)
    add_table(
        doc,
        ["检查项", "通过标准"],
        [
            ["客户侧", "客户能注册/登录、提交需求、入库、出库、维护 SKU、查看账单和账号资料。"],
            ["运营侧", "运营能审核询盘、入库、库存、出库、物流、账单、客户状态和工单。"],
            ["仓库侧", "仓库能扫码定位任务，推进收货、上架、拣货、复核、打包、交运和异常登记。"],
            ["管理侧", "老板能看待办、异常、账单、库存风险、权限和上线体检。"],
            ["中文一致性", "页面文案、下载模板、CSV/Excel 表头和教程内容都以中文为主。"],
        ],
        widths=[2200, 7160],
    )


def main():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = setup_document()
    cover(doc)
    toc_and_map(doc)
    customer_sections(doc)
    ops_sections(doc)
    warehouse_sections(doc)
    flow_sections(doc)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
