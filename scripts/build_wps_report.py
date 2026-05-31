from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT_DIR = ROOT / "deliverables"
OUT_DOCX = OUT_DIR / "英国海外仓仓储系统网站与工作流方案_网站开发版.docx"
DOCX_SCREENSHOT = OUT_DIR / "prototype-docx-screenshot.png"

BLUE = "1F4E79"
NAVY = "12324A"
LIGHT_BLUE = "EAF3FB"
LIGHT_GRAY = "F3F6F9"
BORDER = "D8DEE6"
GREEN = "16A34A"
AMBER = "F59E0B"
RED = "DC2626"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "6") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=120, start=160, bottom=120, end=160) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def style_table(table, header_fill=NAVY, header_font="FFFFFF") -> None:
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for row_idx, row in enumerate(table.rows):
        for cell in row.cells:
            set_cell_border(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.15
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
            if row_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor.from_string(header_font)
            else:
                set_cell_shading(cell, "FFFFFF" if row_idx % 2 else "F8FAFC")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_col_widths(table, widths_cm: list[float]) -> None:
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(width)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths_cm: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    hdr = table.rows[0].cells
    for i, header in enumerate(headers):
        hdr[i].text = header
    for row_data in rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            row[i].text = text
    if widths_cm:
        set_col_widths(table, widths_cm)
    set_repeat_table_header(table.rows[0])
    style_table(table)
    doc.add_paragraph()


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.color.rgb = RGBColor.from_string(NAVY if level == 1 else BLUE)
        run.bold = True


def add_body(doc: Document, text: str, bold_prefix: str | None = None) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.25
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r2 = p.add_run(text[len(bold_prefix) :])
        runs = [r1, r2]
    else:
        runs = [p.add_run(text)]
    for run in runs:
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10.5)
        run.font.color.rgb = RGBColor.from_string("1F2933")


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(3)
        p.paragraph_format.line_spacing = 1.2
        run = p.add_run(item)
        run.font.name = "Microsoft YaHei"
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        run.font.size = Pt(10)


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_BLUE) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, "BFD7EA")
    set_cell_margins(cell, top=180, start=220, bottom=180, end=220)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.name = "Microsoft YaHei"
    r1._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r1.font.size = Pt(11)
    r1.font.color.rgb = RGBColor.from_string(NAVY)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor.from_string("334155")
    p2.paragraph_format.line_spacing = 1.25
    doc.add_paragraph()


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("英国海外仓仓储系统网站与工作流方案 | 2026")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor.from_string("64748B")


def set_document_styles(doc: Document) -> None:
    section = doc.sections[0]
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)
    add_footer(section)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)

    for name, size, color in [
        ("Title", 24, NAVY),
        ("Heading 1", 17, NAVY),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 11, BLUE),
    ]:
        style = styles[name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True


def cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(40)
    r = p.add_run("英国海外仓仓储系统网站与工作流方案")
    r.bold = True
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(26)
    r.font.color.rgb = RGBColor.from_string(NAVY)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("面向中国跨境电商与外贸卖家的市场调研、产品方案、数据模型与原型交付")
    r2.font.name = "Microsoft YaHei"
    r2._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r2.font.size = Pt(13)
    r2.font.color.rgb = RGBColor.from_string("475569")

    doc.add_paragraph()
    add_callout(
        doc,
        "项目定位",
        "建设一套服务中国跨境卖家的英国本地仓配中枢，覆盖入库预报、库存可视化、订单履约、FBA 补仓、退货换标、异常协同与费用对账。",
    )

    meta = [
        ["交付对象", "英国海外仓网站与仓储系统"],
        ["目标客户", "中国跨境电商卖家、外贸公司、Amazon/eBay/TikTok/Temu/Shopify 商家"],
        ["交付内容", "市场调研、竞品参考、MVP、页面结构、数据模型、路线图、可运行原型"],
        ["生成日期", date.today().isoformat()],
    ]
    add_table(doc, ["项目", "说明"], meta, [4.0, 11.5])
    doc.add_page_break()


def toc(doc: Document) -> None:
    add_heading(doc, "目录式总览", 1)
    rows = [
        ["1", "市场调研结论", "判断英国海外仓对中国卖家的真实需求与获客切入点"],
        ["2", "竞品参考", "拆解中国跨境 ERP/海外仓与英国本地 WMS/3PL 的借鉴点"],
        ["3", "MVP 功能清单", "明确第一版商用闭环必须上线的模块"],
        ["4", "页面结构", "规划官网、客户门户、仓库后台的信息架构"],
        ["5", "数据模型", "定义客户、SKU、ASN、库存流水、订单、退货、FBA、账单等实体"],
        ["6", "开发路线图", "安排从原型到 MVP、集成和效率提升的阶段"],
        ["7", "可运行网站原型", "说明当前 Next.js 原型范围、验证结果与文件位置"],
    ]
    add_table(doc, ["序号", "交付项", "内容说明"], rows, [1.4, 4.2, 10.0])


def market_section(doc: Document) -> None:
    add_heading(doc, "1. 市场调研结论", 1)
    add_callout(
        doc,
        "核心结论",
        "英国海外仓不是单纯仓库，而是中国卖家的英国本地运营节点。客户购买的是平台履约稳定性、库存透明、退货再销售能力、FBA 周边服务、中文客服和合规资料留痕。",
    )
    add_body(
        doc,
        "英国电商消费成熟，Amazon、eBay、TikTok Shop、Temu、Shopify 等平台共同推动中国卖家把库存前置到英国。客户诉求已经从“便宜发货”升级为“平台合规 + 本地履约 + 退货处理 + 账单透明”。",
    )
    add_bullets(
        doc,
        [
            "本地履约：迟发、取消、无追踪号会影响平台账号表现。",
            "FBA 周边：补仓、中转、移除接收、退货换标是高频刚需。",
            "多平台库存：同一批英国库存需要服务 Amazon、eBay、TikTok Shop、Temu、Shopify。",
            "退货再销售：退货质检、拍照、换标、重新上架直接影响利润。",
            "合规信任：VAT、EORI、FHDDS、进口清关资料需要可留痕。",
            "中文协作：中国卖家需要中文客服、微信/企微沟通和异常工单。",
        ]
    )
    add_heading(doc, "目标客户分层", 2)
    rows = [
        ["Amazon UK/FBA 卖家", "FBA 补仓、贴 FNSKU、移除接收、退货换标", "FBA 中转仓 + 退货换标 + 补仓预约"],
        ["eBay UK 卖家", "英国本地一件代发、退货地址、追踪号回传", "英国本土发货，提高买家体验"],
        ["TikTok Shop 卖家", "爆单出库、快速发货、售后退货", "支持 TikTok 爆款备货和高峰出库"],
        ["Temu 本地仓卖家", "本地库存、本地尾程、低价 SKU 周转", "英国本地仓，支持平台本地化履约"],
        ["Shopify/DTC 卖家", "品牌包装、独立站订单履约、退货体验", "英国本地履约 + 品牌包装 + 退货体验"],
        ["外贸 B2B 客户", "小批量现货、样品寄送、托盘出库", "英国仓现货，小批量快速交付客户"],
    ]
    add_table(doc, ["客户类型", "核心需求", "获客表达"], rows, [3.3, 6.2, 6.2])
    add_heading(doc, "公开资料参考", 2)
    add_bullets(
        doc,
        [
            "ONS 英国互联网零售数据：https://www.ons.gov.uk/businessindustryandtrade/retailindustry/timeseries/j4mc",
            "GOV.UK FHDDS：https://www.gov.uk/guidance/fulfilment-house-due-diligence-scheme",
            "GOV.UK EORI：https://www.gov.uk/eori",
            "Amazon UK MCF：https://sell.amazon.co.uk/fulfilment-by-amazon/fba-multi-channel",
            "TikTok Shop UK FBT：https://newsroom.tiktok.com/tiktok-shop-uk-fbt-launch?lang=en-GB",
        ]
    )


def competitors_section(doc: Document) -> None:
    add_heading(doc, "2. 竞品参考", 1)
    add_body(doc, "竞品应分成两类看：一类决定中国客户如何理解服务，一类决定英国仓库如何高效执行。")
    rows = [
        ["店小秘 / 马帮", "跨境 ERP / OMS", "学习中国卖家的操作语言：店铺授权、订单、库存、物流、财务"],
        ["易仓 / 通途", "海外仓 WMS / ERP", "学习一件代发、FBA 中转、退货换标、PDA 作业、库位管理"],
        ["4PX / 谷仓 / 万邑通", "跨境物流与海外仓", "学习中文服务包装：头程、清关、海外仓、尾程、退货"],
        ["Mintsoft / Peoplevox", "英国/欧洲 3PL WMS", "学习收货、上架、拣货、打包、客户门户和仓库报表"],
        ["ShipStation / Linnworks", "打单与多渠道订单", "学习英国本地承运商、面单、追踪、OMS 结构"],
        ["OE Tech / YoYoParcel / FixPrep", "英国海外仓 / FBA Prep", "学习 FBA Prep、退货处理、快速周转和信任点表达"],
    ]
    add_table(doc, ["参考对象", "定位", "可借鉴点"], rows, [3.5, 4.2, 8.0])
    add_callout(
        doc,
        "产品启发",
        "中文客户门户参考店小秘、马帮、4PX、万邑通、谷仓；仓库内部后台参考易仓、通途、Mintsoft、Peoplevox；英国本地打单和渠道能力参考 ShipStation、Linnworks。",
    )
    add_heading(doc, "开源项目参考", 2)
    rows2 = [
        ["GreaterWMS", "Python/Django、Quasar/Vue", "参考 ASN、库位、库存、扫码、盘点"],
        ["InvenTree", "Python/Django", "参考 SKU、库存台账、条码、标签"],
        ["OpenBoxes", "Grails / React", "参考库存事务、收发货、库存移动审计"],
        ["OpenShip", "Node / Keystone / Next.js", "参考订单路由、履约伙伴模型"],
        ["Medusa / Vendure", "TypeScript", "参考订单、库存、多渠道和插件化"],
        ["Karrio", "Python", "参考多承运商报价、下单、面单、追踪抽象"],
    ]
    add_table(doc, ["项目", "技术栈", "参考价值"], rows2, [3.2, 4.5, 8.0])
    add_heading(doc, "本轮深度调研补充", 2)
    add_callout(
        doc,
        "流程把关结论",
        "中国卖家最熟悉的不是抽象的“数字化海外仓”，而是“建 SKU -> 做入库预报 -> 贴箱唛/预约 -> 到仓签收 -> 验货上架 -> 订单履约/退货/FBA -> 费用对账”。网站首页必须把这些流程翻译成清晰的下一步动作。",
    )
    rows3 = [
        ["首页第一屏", "突出“先估算费用、提交询盘、创建入库预报、查进度/工作台”五个动作"],
        ["首页服务范围", "参考成熟海外仓服务页的信息架构，把服务拆成到仓前协同、海外仓储、一件代发、尾程配送、FBA 补仓、退货与定制服务"],
        ["客户看板", "从仓库内部大屏改成客户视角：即将到仓、可售库存、今日出库、费用待确认、待补资料"],
        ["入库页", "用 5 步解释入库要求，并把表单改成“基础信息、货物信息、资料确认”三步向导"],
        ["服务页", "每个服务都写清适合谁、客户准备什么、仓库做什么、完成后客户看到什么"],
        ["美术方向", "深英伦蓝、冷灰白、少量金铜点睛；真实仓库/扫码/标签/打包场景，不做泛科技物流图"],
    ]
    add_table(doc, ["模块", "优化方向"], rows3, [3.2, 12.4])


def mvp_section(doc: Document) -> None:
    add_heading(doc, "3. MVP 功能清单", 1)
    add_callout(
        doc,
        "MVP 合格标准",
        "第一版必须让客户下单、仓库能做、库存准确、异常可追踪、费用可预估、账单能收钱。复杂平台 API、自动比价、PDA 原生 App 和高级 BI 可以后置。",
    )
    rows = [
        ["客户注册/KYC", "P0", "公司资料、联系人、VAT、EORI、平台店铺、文件上传、审核"],
        ["SKU 管理", "P0", "中文名、英文名、HS Code、FNSKU、尺寸重量、敏感品属性"],
        ["入库预报 ASN", "P0", "箱数、托盘数、SKU 数量、装箱单、商业发票、预计到仓"],
        ["收货上架", "P0", "实收数量、差异、破损拍照、库位、待上架转可售"],
        ["库存管理", "P0", "可售、待上架、锁定、冻结、异常、在途、库存流水"],
        ["订单履约", "P0", "手工创建、批量导入、库存校验、拣货、打包、出库"],
        ["退货换标", "P0", "RMA、退货追踪号、质检、拍照、重上架、换标、销毁"],
        ["FBA 中转", "P0", "Shipment ID、FNSKU、箱标、贴标、换箱、打托、出库"],
        ["异常管理", "P0", "入库差异、破损、缺货、地址异常、标签错误、客户确认"],
        ["费用预估", "P0", "平台、服务、订单量、入库量、库存体积、退货和 FBA 处理量的月度费用拆分，并可带入正式询盘"],
        ["账单管理", "P0", "入库费、仓租、操作费、尾程费、退货费、换标费、FBA 费"],
    ]
    add_table(doc, ["模块", "优先级", "MVP 范围"], rows, [3.2, 1.8, 10.6])
    add_heading(doc, "后置能力", 2)
    add_bullets(
        doc,
        [
            "P1：批量导入/导出增强、邮件通知、低库存提醒、客户报价模板、仓库作业看板、单渠道面单 API 试点。",
            "P2：Amazon/eBay/Shopify/TikTok API、多仓管理、PDA/PWA、自动物流比价、客户钱包、BI 报表、财务系统对接。",
        ]
    )


def page_structure_section(doc: Document) -> None:
    add_heading(doc, "4. 页面结构", 1)
    add_body(doc, "页面结构分成官网、客户门户和仓库后台。官网负责获客与信任，客户门户负责客户提交和查询，仓库后台负责执行、校验、计费和留痕。")
    rows = [
        ["官网", "首页、服务能力、平台方案、合规中心、仓储系统、价格说明、费用预估器、帮助中心、联系我们"],
        ["客户门户", "工作台、账户认证、SKU、入库、库存、订单、退货、FBA、异常、账单、渠道、帮助"],
        ["仓库后台", "运营工作台、客户管理、商品管理、入库作业、库存作业、拣货波次、打包复核、退货质检、FBA 作业、异常、费用计费、系统设置"],
    ]
    add_table(doc, ["端口", "页面结构"], rows, [3.0, 12.6])
    add_heading(doc, "客户门户工作台指标", 2)
    add_bullets(
        doc,
        [
            "可售库存 SKU 数、待入库箱数、今日待出库订单、异常订单、本月预估费用、低库存 SKU。",
            "待办事项：待提交箱唛、待确认入库差异、待处理退货质检、待支付账单、待补充申报信息。",
            "订单概览：待审核、待拣货、已出库、已妥投、派送异常。",
        ]
    )
    add_heading(doc, "仓库后台工作台指标", 2)
    add_bullets(
        doc,
        [
            "今日预计到仓 ASN、待卸货箱数、待上架 SKU、待拣订单、待打包订单、待交接包裹。",
            "作业队列：入库预约、上架任务、拣货波次、退货质检、异常复核。",
            "效率指标：入库完成率、按时出库率、拣货错误率、退货质检平均时长、仓位利用率。",
        ]
    )


def data_model_section(doc: Document) -> None:
    add_heading(doc, "5. 数据模型", 1)
    add_callout(
        doc,
        "设计原则",
        "所有业务单据必须带 customer_id；所有库存变动必须写 InventoryLedger；所有计费必须来自 ChargeEvent；合规字段变更必须写 ComplianceAuditLog；重要 API 必须支持 idempotency_key。",
    )
    rows = [
        ["CustomerAccount", "客户主体、联系人、状态、风险等级"],
        ["CustomerComplianceProfile", "VAT、EORI、平台账号、KYC、FHDDS 尽调状态"],
        ["Sku", "SKU、FNSKU、ASIN、HS Code、申报资料、尺寸重量、敏感属性"],
        ["InboundASN / ASNLine", "入库预报、运输方式、ETA、箱数、预报数量、实收差异"],
        ["InventoryBalance", "按客户、仓库、SKU、库位、状态聚合的库存快照"],
        ["InventoryLedger", "库存流水，记录入库、锁定、出库、退货、调整、销毁"],
        ["SalesOrder / Shipment", "订单、平台单号、地址、拣货、打包、承运商、追踪号"],
        ["ReturnOrder / RelabelTask", "退货、质检、换标、重上架、销毁、转寄"],
        ["FbaTransfer", "FBA Shipment ID、FNSKU、箱标、贴标、打托和出库"],
        ["ChargeEvent / Invoice", "计费事件、月结账单、税费、付款状态"],
        ["Ticket", "异常工单、客户确认、内部协作、SLA"],
    ]
    add_table(doc, ["实体", "用途"], rows, [4.2, 11.4])
    add_heading(doc, "关键状态机", 2)
    rows2 = [
        ["ASN", "draft -> submitted -> arrived -> receiving -> putaway_completed -> closed"],
        ["订单", "created -> allocated -> picking -> packed -> shipped -> delivered -> closed"],
        ["退货", "created -> received -> inspected -> disposition_pending -> completed"],
        ["FBA", "draft -> allocated -> prep_in_progress -> packed -> shipped -> closed"],
        ["账单", "draft -> issued -> partially_paid / paid / overdue / void"],
    ]
    add_table(doc, ["业务对象", "状态流转"], rows2, [3.0, 12.6])


def roadmap_section(doc: Document) -> None:
    add_heading(doc, "6. 开发路线图", 1)
    rows = [
        ["阶段 1", "产品蓝图与原型", "市场调研、竞品、PRD、数据模型、路线图、可运行原型", "已完成"],
        ["阶段 2", "MVP 基础工程", "登录权限、客户/KYC、SKU、ASN、收货上架、库存流水", "下一步"],
        ["阶段 3", "履约闭环", "订单创建/导入、库存锁定、拣货、打包、出库、追踪号", "计划中"],
        ["阶段 4", "退货/FBA/异常", "退货质检、换标、FBA 中转、异常中心、客户确认流", "计划中"],
        ["阶段 5", "计费与账单", "费率卡、计费事件、月结账单、付款凭证、财务确认", "计划中"],
        ["阶段 6", "集成与效率提升", "平台 API、面单 API、企业微信通知、PDA/PWA、多承运商规则", "计划中"],
    ]
    add_table(doc, ["阶段", "主题", "交付内容", "状态"], rows, [2.0, 3.2, 8.2, 2.2])
    add_heading(doc, "主要风险与应对", 2)
    rows2 = [
        ["库存账实不一致", "所有库存变化写流水，余额可由流水重建"],
        ["账单争议", "每笔费用关联来源单据和操作日志"],
        ["合规资料不完整", "KYC、VAT、EORI、FHDDS 审核状态前置"],
        ["平台 API 复杂", "第一版用 CSV/Excel，API 分阶段接入"],
        ["英国快递面单差异大", "先手动/半自动，选单渠道试点"],
        ["退货识别困难", "支持 tracking、原订单、收件人、SKU 模糊匹配"],
    ]
    add_table(doc, ["风险", "应对策略"], rows2, [5.0, 10.6])


def prototype_section(doc: Document) -> None:
    add_heading(doc, "7. 可运行网站原型", 1)
    add_body(doc, "当前已完成一个 Next.js + React + Tailwind 的可运行原型，包含官网入口、正式营销页面、费用预估到询盘交接、询盘/入库表单、查进度、客户门户、仓库后台、数据模型和路线图。")
    rows = [
        ["项目路径", r"E:\仓储系统\warehouse-system"],
        ["本地地址", "http://localhost:3000"],
        ["主要页面", "官网入口、首页行动入口、首页服务范围矩阵、服务/流程/价格/帮助/联系、费用预估器、预估带入询价、报价跟进工作流、动态客户门户、询盘表单、入库预报、查进度、补交资料、仓库后台、数据模型、路线图"],
        ["验证结果", "lint 通过；production build 通过；首页首屏已突出先估算月成本、直接报价、入库预报和查进度；首页新增“服务包含什么”矩阵，覆盖到仓前协同、英国仓储、一件代发、尾程配送、FBA 补仓、退货与定制服务，并为每张卡提供具体下一步；费用预估器已前置到价格页；询价成功页已增加复制编号和下一步说明；运营后台已支持询盘状态、正式报价草稿、报价有效期和客服备注；客户查进度可看到报价状态；客户门户可显示报价待确认；补交资料 API 已支持真实文件上传到本地目录；390px 移动端无横向滚动"],
    ]
    add_table(doc, ["项目", "说明"], rows, [3.5, 12.1])

    screenshot = DOCX_SCREENSHOT if DOCX_SCREENSHOT.exists() else ROOT / "prototype-home.png"
    if screenshot.exists():
        add_heading(doc, "原型首页截图", 2)
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(screenshot), width=Inches(4.45))
        caption = doc.add_paragraph("图：英国海外仓履约系统原型首页")
        caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in caption.runs:
            r.font.name = "Microsoft YaHei"
            r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string("64748B")

    add_heading(doc, "建议下一步", 2)
    add_bullets(
        doc,
        [
            "把原型升级成真实 MVP 工程，优先实现客户/KYC、SKU、ASN、库存流水、订单出库。",
            "把动态客户门户接入登录和数据库，实现客户数据隔离；把补交资料升级为对象存储、客服审核、文件权限和仓库任务提醒。",
            "确认英国仓实际报价规则，包括入库、仓储、出库、包材、退货、换标、FBA Prep、尾程附加费。",
            "确认第一批试仓客户的真实 SKU、箱单、订单模板和退货流程，用真实数据校正字段。",
        ]
    )


def build() -> None:
    OUT_DIR.mkdir(exist_ok=True)
    prepare_docx_screenshot()
    doc = Document()
    set_document_styles(doc)
    cover(doc)
    toc(doc)
    market_section(doc)
    competitors_section(doc)
    mvp_section(doc)
    page_structure_section(doc)
    data_model_section(doc)
    roadmap_section(doc)
    prototype_section(doc)

    doc.save(OUT_DOCX)
    print(OUT_DOCX)


def prepare_docx_screenshot() -> None:
    source = ROOT / "prototype-home-service-scope-desktop.png"
    if not source.exists():
        source = ROOT / "prototype-home-stage-entry-desktop.png"
    if not source.exists():
        source = ROOT / "prototype-home-action-entry-desktop.png"
    if not source.exists():
        source = ROOT / "prototype-home.png"
    if not source.exists():
        return
    try:
        from PIL import Image
    except Exception:
        return
    image = Image.open(source)
    width, height = image.size
    # Keep the recognisable first-screen area and avoid making the DOCX page too tall.
    crop_height = min(height, int(width * 0.78))
    cropped = image.crop((0, 0, width, crop_height))
    cropped.save(DOCX_SCREENSHOT)


if __name__ == "__main__":
    build()
