from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "deliverables" / "英国驿站仓储物流工作台操作教程.docx"

BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
INK = RGBColor(15, 23, 42)
MUTED = RGBColor(71, 85, 105)
LIGHT_BLUE = "E8EEF5"
LIGHT_GRAY = "F8FAFC"
BORDER = "CBD5E1"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "4") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), "120")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(0, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            width = widths[min(idx, len(widths) - 1)]
            tc_w = cell._tc.get_or_add_tcPr().find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                cell._tc.get_or_add_tcPr().append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(width))
            cell.width = width
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_margins(cell)
            set_cell_border(cell)


def set_run_font(run, size=None, bold=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = color


def add_para(doc, text="", style=None, size=11, bold=False, color=INK, after=6, before=0):
    p = doc.add_paragraph(style=style)
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.25
    if text:
        run = p.add_run(text)
        set_run_font(run, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_heading("", level=level)
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(10)
        size = 16
        color = BLUE
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(7)
        size = 13
        color = BLUE
    else:
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(5)
        size = 12
        color = DARK_BLUE
    run = p.add_run(text)
    set_run_font(run, size=size, bold=True, color=color)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.25
    run = p.add_run(text)
    set_run_font(run, size=11, color=INK)
    return p


def add_callout(doc, title, body, fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [9360])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    set_run_font(r, size=11, bold=True, color=DARK_BLUE)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.25
    r2 = p2.add_run(body)
    set_run_font(r2, size=10.5, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        set_cell_shading(hdr[i], LIGHT_BLUE)
        p = hdr[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(text)
        set_run_font(run, size=10.5, bold=True, color=DARK_BLUE)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            if i < len(cells):
                set_cell_shading(cells[i], WHITE)
                p = cells[i].paragraphs[0]
                p.paragraph_format.space_after = Pt(0)
                p.paragraph_format.line_spacing = 1.2
                if i == 0 and len(str(text)) <= 8:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(str(text))
                set_run_font(run, size=10, color=INK, bold=(i == 0))
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def apply_styles(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for margin in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, margin, Inches(1))
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    header = section.header.paragraphs[0]
    header.text = "英国驿站仓储/物流工作台操作教程"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    set_run_font(header.runs[0], size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.text = "内部培训资料 - 请以系统最新页面为准"
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(footer.runs[0], size=9, color=MUTED)


def add_cover(doc):
    add_para(doc, "操作手册", size=12, bold=True, color=BLUE, after=6)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("英国驿站仓储/物流工作台操作教程")
    set_run_font(r, size=26, bold=True, color=INK)
    add_para(doc, "面向客户、运营、仓库员工和管理者的日常使用说明", size=13, color=MUTED, after=18)
    add_table(
        doc,
        ["项目", "说明"],
        [
            ["适用对象", "客户、客服/运营、仓库员工、管理员、财务"],
            ["适用系统", "客户工作台、运营后台、仓库作业台、官网入口"],
            ["当前版本", f"V1.0 / {date.today().isoformat()}"],
            ["阅读方式", "先看角色地图，再按自己的角色阅读对应章节。"],
        ],
        [1800, 7560],
    )
    add_callout(
        doc,
        "一句话理解",
        "客户在工作台提交需求、入库、SKU、出库、退货和账单动作；运营在后台审核与推进；仓库在作业台扫码、收货、上架、拣货、打包和交运；老板可以看待办、异常、账单和整体运营状态。",
        fill="EEF6FF",
    )
    doc.add_page_break()


def add_contents(doc):
    add_heading(doc, "目录", 1)
    for item in [
        "1. 系统入口和角色地图",
        "2. 客户工作台操作教程",
        "3. 运营后台操作教程",
        "4. 仓库作业台操作教程",
        "5. 老板/管理者怎么看数据",
        "6. 常见业务流程",
        "7. 常见问题与操作规范",
        "8. 上线前后注意事项",
    ]:
        add_bullet(doc, item)
    doc.add_page_break()


def build_doc():
    doc = Document()
    apply_styles(doc)
    add_cover(doc)
    add_contents(doc)

    add_heading(doc, "1. 系统入口和角色地图", 1)
    add_table(
        doc,
        ["入口", "谁使用", "主要用途"],
        [
            ["官网", "新客户、未登录客户", "了解服务范围、费用结构、提交需求、进入客户工作台。"],
            ["客户工作台", "已注册客户", "查看自己的需求、入库、SKU、库存、出库、物流、账单、退货和工单。"],
            ["运营后台", "客服、运营、管理员、财务", "审核客户、处理需求报价、入库、库存、出库、物流异常、账单和工单。"],
            ["仓库作业台", "仓库员工", "扫码收货、上架、库位管理、拣货、打包复核、交运和异常处理。"],
        ],
        [1900, 2100, 5360],
    )
    add_callout(doc, "账号说明", "客户可以自主注册客户工作台账号；员工必须先加入员工白名单。账号密码不建议写入培训手册，请由管理员单独发放和定期更换。")

    add_heading(doc, "2. 客户工作台操作教程", 1)
    add_para(doc, "客户工作台是客户登录后每天使用的页面。客户只会看到自己账号下的数据，不能看到其他客户的数据。", size=11)
    add_heading(doc, "2.1 注册和登录", 2)
    for step in [
        "打开客户工作台登录页。",
        "没有账号时选择免费注册，填写公司/店铺名称、联系人、手机/微信、邮箱和密码。",
        "注册成功后系统会自动生成客户编号，并直接进入客户工作台。",
        "以后使用注册的手机或邮箱加密码登录。",
        "如果账号被暂停，请联系运营确认原因。"
    ]:
        add_number(doc, step)
    add_heading(doc, "2.2 首页先看什么", 2)
    add_table(
        doc,
        ["模块", "客户应该怎么用"],
        [
            ["账号状态", "确认是否未认证、已认证或暂停。未认证时先补公司、VAT、EORI 和平台店铺信息。"],
            ["我的待办", "优先处理报价确认、入库资料补充、账单确认、物流异常和库存风险。"],
            ["常用功能", "提交需求、入库预报、SKU 档案、出库申请、退货预报、查进度、费用账单、账号资料。"],
            ["自助资料", "下载库存报表、出库明细、费用明细、面单、签收证明、物流异常与赔付资料。"],
        ],
        [1900, 7460],
    )
    add_heading(doc, "2.3 提交需求", 2)
    for step in [
        "点击“提交需求”。",
        "填写销售平台、货量、SKU 数、服务需求、尾程/FBA/退货等信息。",
        "提交后运营会在后台评估，并给出报价或补充问题。",
        "客户可在工作台或账单/报价区域查看报价状态。",
        "确认报价后，可以继续创建入库预报。"
    ]:
        add_number(doc, step)
    add_heading(doc, "2.4 入库预报", 2)
    add_table(
        doc,
        ["需要准备", "说明"],
        [
            ["预计到仓日期", "不要早于当天，尽量按真实送仓时间填写。"],
            ["运输方式/追踪号", "快递、卡车、海运、空运等；有追踪号时尽量填写。"],
            ["箱数和 SKU 数", "用于仓库提前安排收货和验货。"],
            ["SKU 明细", "建议按“SKU、品名、数量、箱数”整理。"],
            ["附件资料", "装箱单、外箱标签、FBA 标签、产品图片等。"],
        ],
        [2100, 7260],
    )
    add_heading(doc, "2.5 SKU 档案", 2)
    for step in [
        "先维护 SKU，再进行入库、库存预警、出库和费用核算。",
        "单个新增时填写 SKU 编码、商品名称、条码、分类和预警库存。",
        "批量新增时下载 SKU 模板，按中文表头填写后导入。",
        "不要随意修改已经有库存流水的 SKU 编码；如确实要改，请联系运营。"
    ]:
        add_bullet(doc, step)
    add_heading(doc, "2.6 出库申请", 2)
    for step in [
        "进入“出库申请”。",
        "选择已维护的 SKU，填写数量、物流渠道、收件人、地址和发货要求。",
        "提交后系统会预占库存，避免同一库存被重复出库。",
        "运营审核库存、地址、渠道、费用后，仓库开始拣货、打包和交运。",
        "客户可以在出库页面和查进度页面查看面单、物流状态、异常和赔付进度。"
    ]:
        add_number(doc, step)
    add_heading(doc, "2.7 退货预报", 2)
    for step in [
        "进入“退货预报”。",
        "填写平台订单号、买家退货追踪号、退货原因、SKU 和处理偏好。",
        "仓库到货后会做质检，运营更新结果。",
        "可售商品重新上架；不可售商品按维修、报废、转寄或重发处理。"
    ]:
        add_number(doc, step)
    add_heading(doc, "2.8 费用账单", 2)
    add_para(doc, "客户在账单页查看正式账单、报价方案、付款状态和费用差异。账单有问题时不要私下口头处理，应在系统里提交争议或补充付款凭证，方便后续核对。")
    add_table(
        doc,
        ["状态", "客户动作"],
        [
            ["待确认", "核对费用明细，确认无误后点击确认。"],
            ["有争议", "填写争议说明，例如费用项、数量、订单号或截图依据。"],
            ["付款待复核", "提交付款参考号或上传付款凭证。"],
            ["已支付", "归档留存，后续可下载对账资料。"],
        ],
        [1800, 7560],
    )
    add_heading(doc, "2.9 工单和异常", 2)
    add_para(doc, "遇到物流异常、库存调整、账单争议、退货售后或资料补充时，优先在工作台提交工单。工单应写清楚单号、SKU、问题、期望处理方式和附件。")

    add_heading(doc, "3. 运营后台操作教程", 1)
    add_para(doc, "运营后台是内部人员处理客户业务的主入口。运营的核心目标是让客户看到清楚的下一步，让仓库拿到完整资料，让财务有可核对记录。")
    add_heading(doc, "3.1 运营每天先看什么", 2)
    for item in [
        "运营待办：新需求、报价待确认、入库资料缺失、出库异常、物流异常、账单待复核。",
        "客户认证：查看新注册客户资料是否完整，必要时更新为已认证或暂停。",
        "入库队列：检查预计到仓、追踪号、箱数、SKU 明细和附件是否完整。",
        "出库队列：检查库存、地址、渠道、面单、拣货状态和异常阻塞。",
        "账单队列：核对费用规则、月结账单、付款凭证和争议处理。"
    ]:
        add_bullet(doc, item)
    add_heading(doc, "3.2 客户认证", 2)
    for step in [
        "进入运营后台客户认证区域。",
        "查看公司名称、联系人、手机、邮箱、VAT、EORI、平台店铺和经营地址。",
        "资料完整且确认合作后，将账号状态改为“已认证”。",
        "资料缺失时保留“未认证”，并让客户补充资料。",
        "风险客户或暂停合作客户改为“暂停”，并记录原因。"
    ]:
        add_number(doc, step)
    add_heading(doc, "3.3 报价和需求处理", 2)
    add_table(
        doc,
        ["步骤", "操作重点"],
        [
            ["新需求", "先判断服务范围、货量、SKU、品类、尾程、FBA 和退货需求。"],
            ["联系客户", "状态更新为已联系，并记录跟进备注。"],
            ["录入报价", "填写入库费、仓储费、出库费、退货费、FBA 费、增值服务费和有效期。"],
            ["客户确认", "客户确认后进入后续入库或正式合作流程。"],
        ],
        [1800, 7560],
    )
    add_heading(doc, "3.4 入库审核", 2)
    for step in [
        "检查 ASN 是否有预计到仓日期、运输方式、追踪号、箱数、SKU 数和附件。",
        "资料完整时推进到资料审核通过或预约确认。",
        "资料缺失时标记待补资料，并通过待办或工单提醒客户。",
        "到仓后跟进仓库收货、差异、上架状态。",
        "入库推进到已上架后，系统会写入库存余额和库存流水。"
    ]:
        add_number(doc, step)
    add_heading(doc, "3.5 出库和物流处理", 2)
    for step in [
        "审核出库申请中的 SKU、数量、库存、收件地址和物流渠道。",
        "按规则试算运费或选择承运商服务。",
        "生成面单或追踪号后，同步给客户侧。",
        "回传仓库处理中、已交接承运商、运输途中、派送中、已签收或异常节点。",
        "遇到派送失败、改派、签收证明或赔付时，在物流异常区登记处理。"
    ]:
        add_number(doc, step)
    add_heading(doc, "3.6 账单和费用", 2)
    add_para(doc, "运营需要保证费用来源能追溯到业务单据。生成账单前先核对客户、费用规则、关联单据、数量、金额和到期日。客户提交付款或争议后，运营在后台复核并更新状态。")

    add_heading(doc, "4. 仓库作业台操作教程", 1)
    add_para(doc, "仓库作业台面向实际操作人员，重点是收货、上架、扫码、拣货、打包、交运和异常处理。")
    add_heading(doc, "4.1 入库收货与上架", 2)
    for step in [
        "打开仓库作业台，查看“入库收货与上架”任务。",
        "按 ASN、追踪号或客户信息找到对应任务。",
        "核对箱数、SKU、装箱单、外箱标签和实际到货情况。",
        "有差异时登记收货差异，例如少收、多收、破损、SKU 不符、缺资料。",
        "确认无误后推进状态：已到仓 -> 收货验收中 -> 已收货 -> 已上架。",
        "上架时填写库位，系统会写入库存余额和库存流水。"
    ]:
        add_number(doc, step)
    add_heading(doc, "4.2 扫码怎么用", 2)
    add_table(
        doc,
        ["扫码对象", "系统会做什么"],
        [
            ["ASN/入库单", "定位入库任务，查看客户、箱数、SKU 和资料状态。"],
            ["出库单/波次/拣货单", "定位出库任务，进入拣货、配货、打包或交运流程。"],
            ["SKU 条码", "核对商品是否属于当前任务，减少错拣。"],
            ["库位码", "确认从哪个库位拣货或上架到哪个库位。"],
            ["追踪号/面单", "定位物流任务，核对交运和签出信息。"],
        ],
        [2100, 7260],
    )
    add_heading(doc, "4.3 出库拣货、打包和交运", 2)
    for step in [
        "查看“出库下架、拣货、打包与签出”任务。",
        "选择拣货模式：集中分拣、拣货车分拣或按单分拣。",
        "打印拣货单或批量拣货单。",
        "按库位和 SKU 扫码拣货。",
        "打包复核时逐个扫描 SKU，确认数量和订单一致。",
        "称重后签出，推进到待交运或已发货。",
        "如果发现缺货、错货、破损、面单错误或截单，登记异常。"
    ]:
        add_number(doc, step)
    add_heading(doc, "4.4 库位管理", 2)
    add_para(doc, "仓库可维护库区、库位、容量、库位类型和是否允许混放。建议真实仓库先给每个货架、层、格口贴库位码，再在系统中导入库位。")
    add_bullet(doc, "库位命名建议：库区-货架-层-格，例如 A-01-04。")
    add_bullet(doc, "退货待检、残次品、冻结库存建议使用单独库区。")
    add_bullet(doc, "不允许混放的库位不要把不同客户或不同 SKU 放在一起。")

    add_heading(doc, "5. 老板/管理者怎么看数据", 1)
    add_para(doc, "管理者不需要处理每一张单，但要每天看风险和效率。")
    add_table(
        doc,
        ["看什么", "怎么看"],
        [
            ["待办数量", "判断今天是否有客户等待、仓库阻塞或财务待复核。"],
            ["物流异常", "重点看派送失败、改派、签收证明、赔付金额和处理状态。"],
            ["库存风险", "看低库存、库龄偏高、冻结库存、残次品和盘点差异。"],
            ["账单金额", "看待确认、待支付、争议和已支付金额。"],
            ["仓库效率", "看入库是否及时上架，出库是否及时拣货和交运。"],
        ],
        [2000, 7360],
    )

    add_heading(doc, "6. 常见业务流程", 1)
    add_heading(doc, "6.1 新客户从注册到入库", 2)
    for step in [
        "客户注册客户工作台账号。",
        "客户完善公司、VAT、EORI 和平台店铺资料。",
        "运营审核客户资料并更新认证状态。",
        "客户提交需求或直接创建入库预报。",
        "运营审核资料，仓库收货上架，客户看到库存更新。"
    ]:
        add_number(doc, step)
    add_heading(doc, "6.2 一票出库从客户提交到签收", 2)
    for step in [
        "客户先维护 SKU，再提交出库申请。",
        "系统预占库存。",
        "运营审核地址、渠道、费用和面单。",
        "仓库打印拣货单，扫码拣货并打包复核。",
        "运营或仓库回传物流节点。",
        "客户在工作台查看最新物流、签收证明或异常处理结果。"
    ]:
        add_number(doc, step)
    add_heading(doc, "6.3 物流异常处理", 2)
    for step in [
        "发现异常后，运营在物流异常模块记录异常类型和说明。",
        "如果需要客户补充信息，客户在工作台提交工单。",
        "需要改派时记录改派地址和要求。",
        "需要赔付时记录金额、状态和证明材料。",
        "处理完成后更新状态，让客户侧同步看到结果。"
    ]:
        add_number(doc, step)

    add_heading(doc, "7. 常见问题与操作规范", 1)
    add_table(
        doc,
        ["问题", "处理方式"],
        [
            ["客户看不到数据", "确认客户是否使用正确账号登录，业务单据是否归属到该客户编号。"],
            ["入库资料不完整", "让客户补装箱单、标签、追踪号或 SKU 明细，不要口头补资料。"],
            ["SKU 不存在", "先让客户在 SKU 档案中建档，再提交入库或出库。"],
            ["库存不够出库", "运营确认是否需要库存调整、补货、拆单或取消部分出库。"],
            ["面单或地址有问题", "先阻塞出库，确认地址、渠道和费用后再继续。"],
            ["账单有争议", "客户在系统提交争议说明，运营复核后更新状态。"],
        ],
        [2200, 7160],
    )
    add_callout(doc, "操作原则", "系统里能记录的事情，不要只在微信里说。单号、SKU、数量、地址、费用、异常和处理结果都要回到系统中留痕。", fill="FFF7ED")

    add_heading(doc, "8. 上线前后注意事项", 1)
    add_bullet(doc, "员工账号必须走白名单，不建议共享账号。")
    add_bullet(doc, "客户密码和员工密码不要写进培训手册，单独发放。")
    add_bullet(doc, "客户可见内容、下载模板、CSV/Excel 表头都应以中文为主。")
    add_bullet(doc, "真实承运商 API、正式面单、对象存储、短信/邮件通知等能力上线前，需要再做一次专项验收。")
    add_bullet(doc, "每次发版后至少检查客户登录、运营登录、入库、出库、账单、扫码和导出。")

    add_heading(doc, "附录：角色速查表", 1)
    add_table(
        doc,
        ["角色", "每天最常用", "最容易出错的点"],
        [
            ["客户", "提交需求、入库、SKU、出库、退货、账单、工单", "忘记先建 SKU、资料只发微信不上传系统。"],
            ["运营", "客户认证、报价、入库审核、物流异常、账单复核", "状态不更新、异常没有写清楚下一步。"],
            ["仓库", "扫码收货、上架、拣货、打包、交运、异常登记", "不扫 SKU 或库位，导致库存和实物不一致。"],
            ["财务", "账单确认、付款凭证、争议处理、月结导出", "费用来源没有关联单据。"],
            ["管理员", "员工白名单、权限、客户状态、上线体检", "账号共享或权限过宽。"],
        ],
        [1400, 4200, 3760],
    )

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    return OUT


if __name__ == "__main__":
    print(build_doc())
